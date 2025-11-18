import os
import re
import io
import zipfile
import tempfile
import traceback
import base64
from typing import Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
import requests

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# Optional image/pdf libs
try:
    from pdf2image import convert_from_path
except Exception:
    convert_from_path = None

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import pytesseract
except Exception:
    pytesseract = None

# Optional Gemini
try:
    import google.generativeai as genai
except Exception:
    genai = None


# -------------------------------
# CONFIGURATION
# -------------------------------
OCR_API_KEY = os.getenv("OCR_API_KEY", "K85450490888957").strip() or None
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyClP2B1jdADvbxd8I96w5Fok8aZZQfXEbQ").strip() or None
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-pro-latest")

if genai and GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
    except Exception:
        genai = None

app = FastAPI(title="Doc Redactor API", version="2.0")


# -------------------------------
# PII LABEL LIST
# -------------------------------
PII_LABELS = [

    # -----------------------------
    # 1. Government Issued ID
    # -----------------------------
    "government issued id", "Government Issued ID", "GOVERNMENT ISSUED ID",
    "govt issued id", "gov issued id", "gov issued identification",
    "gov id", "govt id", "government id", "government identification",
    "id issued by government", "government identity card",
    "id card", "identity card", "identification id",
    "official id", "official identification", "national id",
    "national identification", "gov identity",

    # -----------------------------
    # 2. Social Security Number
    # -----------------------------
    "social security number", "Social Security Number", "SOCIAL SECURITY NUMBER",
    "ssn", "SSN", "S.S.N.", "social security no", "ss number",
    "soc sec no", "ssn number", "social sec number", "social security #",

    # -----------------------------
    # 3. Tax ID
    # -----------------------------
    "tax id", "Tax ID", "TAX ID", "tax identification number",
    "tin", "TIN", "T.I.N.", "tax no", "tax number",
    "taxpayer id", "tax payer number",

    # -----------------------------
    # 4. Federal Employer ID
    # -----------------------------
    "federal employer id", "Federal Employer ID", "FEDERAL EMPLOYER ID",
    "employer id", "employer identification", "feid", "FEID", "F.E.I.D.",

    # -----------------------------
    # 5. FEIN
    # -----------------------------
    "fein", "FEIN", "F.E.I.N.", "federal employer identification number",
    "fein number", "federal ein", "employer ein",

    # -----------------------------
    # 6. Driver's License
    # -----------------------------
    "driver's license", "Driver's License", "Driver' s License","License","DRIVER'S LICENSE",
    "drivers license", "driver license", "driving license",
    "dl number", "DL", "D.L.", "license number", "driver id",

    # -----------------------------
    # 7. Identification Card
    # -----------------------------
    "identification card", "Identification Card", "ID card",
    "identity card", "id", "ID", "identification", "id number",
    "identification number",

    # -----------------------------
    # 8. Passport
    # -----------------------------
    "passport", "Passport", "PASSPORT", "passport number",
    "passport no", "pp number", "passport id",

    # -----------------------------
    # 9. Military ID
    # -----------------------------
    "military id", "Military ID", "MILITARY ID",
    "army id", "navy id", "airforce id", "defense id",
    "military identification",

    # -----------------------------
    # 10. Date of Birth
    # -----------------------------
    "date of birth", "Date of Birth", "DATE OF BIRTH",
    "dob", "DOB", "birth date", "birth info","D.o.B.","DOB",
    "date born", "born on", "birthdate","D.O.B.",

    # -----------------------------
    # 11. Home Address
    # -----------------------------
    "home address", "Home Address", "HOME ADDRESS",
    "residential address", "residence address", "address", "addr","ADDRESS",
    "street address", "street addr", "residential addr","Address",

    # -----------------------------
    # 12. Home Telephone Number
    # -----------------------------
    "home telephone number", "Home Telephone number",
    "HOME TELEPHONE NUMBER", "telephone number",
    "home phone", "landline", "tel number",

    # -----------------------------
    # 13. Cell Phone Number
    # -----------------------------
    "cell phone number", "Cell phone number", "CELL PHONE NUMBER",
    "mobile number", "mobile no", "cell number", "phone number",
    "contact number", "contact no","ph number","Cell No",

    # -----------------------------
    # 14. Email Address
    # -----------------------------
    "email address", "Email Address", "EMAIL ADDRESS",
    "email", "e-mail", "email id", "mail id","Email","email ID","eMail","gmail","g-mail",

    # -----------------------------
    # 15. Social Media Contact Information
    # -----------------------------
    "social media contact information", "Social Media Contact Information",
    "SOCIAL MEDIA CONTACT INFORMATION", "social media info",
    "social handle", "social contact", "social media account",

    # -----------------------------
    # 16. Health Insurance Policy Number
    # -----------------------------
    "health insurance policy number", "Health Insurance Policy Number",
    "insurance policy number", "policy number", "policy no",
    "health insurance number", "insurance number",

    # -----------------------------
    # 17. Medical Record Number
    # -----------------------------
    "medical record number", "Medical Record Number",
    "MRN", "mrn", "medical record no", "med record number","medical","record","number",

    # -----------------------------
    # 18. Claim Number
    # -----------------------------
    "claim number", "Claim Number", "CLAIM NUMBER",
    "claim no", "claim id",

    # -----------------------------
    # 19. Patient Account Number
    # -----------------------------
    "patient account number", "Patient Account Number",
    "PATIENT ACCOUNT NUMBER", "patient id", "patient account",

    # -----------------------------
    # 20. File Number
    # -----------------------------
    "file number", "File Number", "FILE NUMBER",
    "file no", "file id", "file reference",

    # -----------------------------
    # 21. Chart Number
    # -----------------------------
    "chart number", "Chart Number", "CHART NUMBER",
    "chart no", "chart id",

    # -----------------------------
    # 22. Individual Financial Account Number
    # -----------------------------
    "individual financial account number", "Individual Financial Account Number",
    "financial account number", "financial account", "account number",

    # -----------------------------
    # 23. Bank Account Number
    # -----------------------------
    "bank account number", "Bank Account Number", "BANK ACCOUNT NUMBER",
    "bank no", "account no", "acct number",

    # -----------------------------
    # 24. Financial Information
    # -----------------------------
    "financial information", "Financial Information",
    "FINANCIAL INFORMATION", "financial data", "financial details",

    # -----------------------------
    # 25. Credit Card Number
    # -----------------------------
    "credit card number", "Credit Card Number", "CREDIT CARD NUMBER",
    "credit card", "card number", "cc number", "card no"
]


PATTERNS = {
    "SSN": r"\b\d{3}-\d{2}-\d{4}\b",
    "Phone": r"\b\+?\d{1,3}?[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}\b",
    "Email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
    "CreditCard": r"\b(?:\d{4}[-\s]?){3}\d{4}\b",
    "Date": r"\b(?:0?[1-9]|1[0-2])[\/\-.](?:0?[1-9]|[12]\d|3[01])[\/\-.]\d{4}\b"
}


# -------------------------------
# File Type Detection
# -------------------------------
def detect_filetype(filename: str):
    ext = filename.lower().split(".")[-1]
    if ext in ("jpg","jpeg","png","bmp","tiff","tif"):
        return "image"
    if ext == "pdf":
        return "pdf"
    if ext == "txt":
        return "text"
    if ext == "docx":
        return "docx"
    return None


# -------------------------------
# OCR.space
# -------------------------------
def ocrspace_extract_bytes(file_bytes, filename, language="eng"):
    if not OCR_API_KEY:
        return ""
    try:
        files = {"file": (filename, file_bytes)}
        data = {"apikey": OCR_API_KEY, "language": language, "OCREngine": 2}
        resp = requests.post("https://api.ocr.space/parse/image", files=files, data=data, timeout=120)
        resp.raise_for_status()
        r = resp.json()
        if r.get("IsErroredOnProcessing"):
            return ""
        return "\n".join(x.get("ParsedText", "") for x in r.get("ParsedResults", []))
    except Exception:
        return ""


# -------------------------------
# Local OCR (Tesseract)
# -------------------------------
def pytesseract_extract_bytes(file_bytes, filename):
    if pytesseract is None or Image is None:
        return ""

    try:
        ftype = detect_filetype(filename)

        if ftype == "image":
            img = Image.open(io.BytesIO(file_bytes))
            return pytesseract.image_to_string(img)

        if ftype == "pdf" and convert_from_path:
            with tempfile.NamedTemporaryFile(suffix=".pdf") as tmp:
                tmp.write(file_bytes); tmp.flush()
                imgs = convert_from_path(tmp.name, dpi=200)
            text = ""
            for img in imgs:
                text += pytesseract.image_to_string(img) + "\n"
            return text
    except:
        return ""
    return ""


# -------------------------------
# DOCX EXTRACT
# -------------------------------
def docx_extract_bytes(file_bytes):
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
            tmp.write(file_bytes); tmp.flush()
            doc = Document(tmp.name)
        return "\n".join(p.text for p in doc.paragraphs)
    except:
        return ""


# -------------------------------
# Extract TEXT From any File
# -------------------------------
def extract_text_from_bytes(file_bytes, filename, language="eng"):
    t = detect_filetype(filename)

    if t == "text":
        return file_bytes.decode("utf-8", errors="ignore")
    if t == "docx":
        return docx_extract_bytes(file_bytes)
    if t in ("pdf", "image"):
        txt = ocrspace_extract_bytes(file_bytes, filename, language)
        if txt.strip(): return txt

        txt = pytesseract_extract_bytes(file_bytes, filename)
        if txt.strip(): return txt

    return ""


# -------------------------------
# Gemini Fix (optional)
# -------------------------------
def fix_text_with_gemini(text):
    if genai is None:
        return text
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        prompt = f"Clean OCR text, fix mistakes only:\n\n{text}"
        resp = model.generate_content([prompt])
        return (resp.text or "").strip()
    except:
        return text


# -------------------------------
# Redaction
# -------------------------------
def blackout(s): return "█" * len(s)

def redact_labels(text):
    for label in PII_LABELS:
        p = rf"({re.escape(label)}\s*[:\-–]\s*)([^\n\r]+)"
        text = re.sub(p, lambda m: m.group(1) + blackout(m.group(2)), text, flags=re.I)
    return text

def redact_patterns(text):
    for _, patt in PATTERNS.items():
        text = re.sub(patt, lambda m: blackout(m.group(0)), text, flags=re.I)
    return text

def redact_text_content(text):
    return redact_patterns(redact_labels(text))


# -------------------------------
# Create DOCX
# -------------------------------
def make_docx_bytes(text):
    doc = Document()
    doc.add_heading("Cleaned & Redacted Document", level=1)
    for line in text.splitlines():
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.read()


# -------------------------------
# Create PDF
# -------------------------------
def make_pdf_bytes(text):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf)
    style = getSampleStyleSheet()["Normal"]
    story = [Paragraph(l if l else " ", style) for l in text.splitlines()]
    doc.build(story)
    buf.seek(0)
    return buf.read()


# -------------------------------
# HEALTHCHECK
# -------------------------------
@app.get("/health")
def health():
    return {"status": "ok"}


# =====================================================
# 1️⃣ MAIN ENDPOINT — ServiceNow-Compatible Response
# =====================================================
@app.post("/process")
async def process_file(
    file: UploadFile = File(...),
    use_ai: bool = Form(False),
    language: str = Form("eng")
):
    """
    Returns:
    {
        "filename": "xyz.zip",
        "content_type": "application/zip",
        "file_base64": "<base64>"
    }
    Works 100% with ServiceNow attachment uploads.
    """
    try:
        data = await file.read()

        if not data:
            raise HTTPException(400, "Empty file")

        text = extract_text_from_bytes(data, file.filename, language)
        if not text.strip():
            raise HTTPException(422, "Unable to extract text")

        if use_ai:
            text = fix_text_with_gemini(text)

        redacted = redact_text_content(text)

        # Create output files
        docx_bytes = make_docx_bytes(redacted)
        pdf_bytes = make_pdf_bytes(redacted)

        # ZIP BOTH FILES
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as z:
            base = os.path.splitext(file.filename)[0]
            z.writestr(f"{base}_cleaned_redacted.docx", docx_bytes)
            z.writestr(f"{base}_cleaned_redacted.pdf", pdf_bytes)

        zip_buf.seek(0)
        zip_bytes = zip_buf.read()

        # Convert ZIP → base64 (ServiceNow compatible)
        zip_base64 = base64.b64encode(zip_bytes).decode()

        return JSONResponse({
            "filename": f"{base}_cleaned_redacted.zip",
            "content_type": "application/zip",
            "file_base64": zip_base64
        })

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(500, f"Server error: {str(e)}")
